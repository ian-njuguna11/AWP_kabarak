import pymysql
from docx import Document
from docx.shared import Inches
import pymysql


pymysql.install_as_MySQLdb()
import mysql.connector

from flask_mysqldb import MySQL

# pymysql.install_as_MySQLdb()
import MySQLdb

from flask import jsonify, render_template, redirect, request, url_for, make_response, Response
import requests
import json
from flask_login import (
    current_user,
    login_required,
    login_user,
    logout_user
)

from app import db, login_manager
from app.base import blueprint
from app.base.forms import LoginForm, CreateAccountForm
from app.base.models import User

from app.base.util import verify_pass

mydb = mysql.connector.connect(
    host="localhost",
    user="root",
    password="",
    database="awp"
)

mycursor = mydb.cursor()




@blueprint.route('/')
def route_default():
    return redirect(url_for('base_blueprint.login'))


## Login & Registration

@blueprint.route('/login', methods=['GET', 'POST'])
def login():
    login_form = LoginForm(request.form)
    if 'login' in request.form:

        # read form data
        username = request.form['username']
        password = request.form['password']

        # Locate user
        user = User.query.filter_by(username=username).first()

        # Check the password
        if user and verify_pass(password, user.password):
            login_user(user)
            return redirect(url_for('base_blueprint.route_default'))

        # Something (user or pass) is not ok
        return render_template('accounts/login.html', msg='Wrong user or password', form=login_form)

    if not current_user.is_authenticated:
        return render_template('accounts/login.html',
                               form=login_form)
    return redirect(url_for('home_blueprint.index'))


@blueprint.route('/register', methods=['GET', 'POST'])
def register():
    login_form = LoginForm(request.form)
    create_account_form = CreateAccountForm(request.form)
    if 'register' in request.form:

        username = request.form['username']
        email = request.form['email']

        # Check usename exists
        user = User.query.filter_by(username=username).first()
        if user:
            return render_template('accounts/register.html',
                                   msg='Username already registered',
                                   success=False,
                                   form=create_account_form)

        # Check email exists
        user = User.query.filter_by(email=email).first()
        if user:
            return render_template('accounts/register.html',
                                   msg='Email already registered',
                                   success=False,
                                   form=create_account_form)

        # else we can create the user
        user = User(**request.form)
        db.session.add(user)
        db.session.commit()

        return render_template('accounts/register.html',
                               msg='User created please <a href="/login">login</a>',
                               success=True,
                               form=create_account_form)

    else:
        return render_template('accounts/register.html', form=create_account_form)


@blueprint.route('/community_unit', methods=['GET', 'POST'])
def comm_unit():
    Province = request.form.get('Province')
    District = request.form.get('District')
    Division = request.form.get('Division')

    mydb = mysql.connector.connect(
        host="localhost",
        user="root",
        password="",
        database="awp"
    )

    mycursor = mydb.cursor()

    # mycursor.execute("CREATE TABLE dummy_community_unit (id INT AUTO_INCREMENT PRIMARY KEY, province VARCHAR(255), district VARCHAR(255),division VARCHAR(255))")
    # mycursor.execute("ALTER TABLE dummy_community_unit ADD COLUMN id INT AUTO_INCREMENT PRIMARY KEY")

    sql = "INSERT INTO dummy_community_unit (province, district,division) VALUES (%s, %s,%s)"
    val = (Province, District,Division)
    mycursor.execute(sql, val)
    #
    # mydb.commit()

    return " <h1>data inserted successfully</h1>"

    # dictToSend = {
    #     "index": Province,
    #     "Outpatient":District,
    #     "Constituency": Division
    # }
    # res = requests.post('http://localhost:8080/api_post', json=dictToSend)
    # print(f'response from server:', res.text)
    # dictFromServer = res.json()

    # cu = request.form['Province','District','Division']
    # return jsonify(
    #     name=Province,
    #     District=District,
    #     Division=Division,
    # )
    # for attribute in cu:
    # return f"{cu[0],cu[1],cu[2]}"

    # if request.method == "POST":
    #     cu = request.form["Province","District"]
    #     return redirect(cu)
    # else:
    #     return   render_template('ui-forms.html')


@blueprint.route('/api_cu', methods=['GET', 'POST'])
def get_cu():
    Province = request.form.get('Province')
    District = request.form.get('District')
    Division = request.form.get('Division')

    print(mydb)

@blueprint.route("/consolidated_moh513")
def consolidated_chv():
    mycursor.execute("SELECT * FROM dummy_community_unit")
    myresult = mycursor.fetchall()

    # moh513 = mycursor.fetchall()
    document = Document()
    document.add_heading('Chew Summary', 0)

    p = document.add_paragraph('This tool   is the  monthly summary of the CHEWs efforts and the services carried at the household levels.   The tool is to be filled monthly by the CHEW using the information from the Community Service Log  (at the end of Month) and after six months ( use the updated Household Register). The information collected Measures the CHWs efforts and services carried out at the household levels. It shows the Community Unit Outputs The information captured on the CHEW summary is also replicated to the CHALK BOARD but interpreted on a negative side to trigger Community actions.')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    # document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='Intense Quote')

    document.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )

    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )

    document.add_picture('yn.png', width=Inches(1.25))

    for record in myresult:
        print(record)

    records = (
        {
            record
        }

    )

    # for record in myresult:
    #     print(record)
    #
    # records = (
    #     {
    #         record
    #     }
    #
    # )

    # records = (
    #
    #     (3, '101', 'Spam'),
    #     (7, '422', 'Eggs'),
    #     (4, '631', 'Spam, spam, eggs, and spam')
    #     )

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in myresult:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc

    document.add_page_break()

    document.save('chewsummarry.docx')

    return " <h1>Document successfully Generated</h1> "




@blueprint.route('/logout')
def logout():
    logout_user()
    return redirect(url_for('base_blueprint.login'))


## Errors

@login_manager.unauthorized_handler
def unauthorized_handler():
    return render_template('page-403.html'), 403


@blueprint.errorhandler(403)
def access_forbidden(error):
    return render_template('page-403.html'), 403


@blueprint.errorhandler(404)
def not_found_error(error):
    return render_template('page-404.html'), 404


@blueprint.errorhandler(500)
def internal_error(error):
    return render_template('page-500.html'), 500
